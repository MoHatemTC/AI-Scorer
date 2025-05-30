import json
from docx import Document
from docx.shared import Pt


class DocumentWriterAgent:
    def __init__(self):
        pass

    def run(self,
            skills_file="extracted_skills.json",
            task_details_file="task_details.json",
            generated_tasks_file="generated_tasks.json",
            rubric_file="rubric.json",
            output_file="Task_and_Rubric_Document.docx"):

        with open(skills_file, encoding="utf-8") as f:
            skills = json.load(f)
        with open(task_details_file, encoding="utf-8") as f:
            task_details = json.load(f)
        with open(generated_tasks_file, encoding="utf-8") as f:
            task_statements = json.load(f)
        with open(rubric_file, encoding="utf-8") as f:
            rubrics = json.load(f)

        doc = Document()
        doc.add_heading("AI-Generated Tasks & Rubrics Document", 0)

        doc.add_heading("Extracted Skills", level=1)
        for skill in skills:
            doc.add_paragraph(f"• {skill}")

        doc.add_heading("Task Details", level=1)
        for item in task_details:
            skill = item['skill']
            details = item['details']
            doc.add_heading(skill, level=2)
            doc.add_paragraph(f"Objective: {details['objective']}")
            doc.add_paragraph(f"Tools: {details['tools']}")
            doc.add_paragraph(f"Input: {details['input']}")
            doc.add_paragraph(f"Output: {details['output']}")
            doc.add_paragraph(f"Duration: {details['duration']}")

        doc.add_heading("Task Statements", level=1)
        for task in task_statements:
            doc.add_heading(task['skill'], level=2)
            doc.add_paragraph(f"Task: {task['task']}")
            doc.add_paragraph(f"Difficulty: {task['difficulty']}")

        doc.add_heading("Rubrics", level=1)
        for item in rubrics:
            doc.add_heading(item['skill'], level=2)
            doc.add_paragraph(f"Task: {item['task']}")
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Criterion'
            hdr_cells[1].text = 'Beginner'
            hdr_cells[2].text = 'Intermediate'
            hdr_cells[3].text = 'Advanced'
            for crit in item['rubric']:
                row_cells = table.add_row().cells
                row_cells[0].text = crit['criterion']
                row_cells[1].text = crit['beginner']
                row_cells[2].text = crit['intermediate']
                row_cells[3].text = crit['advanced']

        doc.save(output_file)
        print(f"Document saved to {output_file}")